# -----------------------------------------------------------------------------
# Title: Distinguished Security Engineer CV
# Author: Ikwuka Okoye (@gadcode)
# Date: October 15, 2023
# Usage: This program uses the Python programming language's docx module to
#        create a professional curriculum vitae (CV).
# -----------------------------------------------------------------------------

from docx import Document


def create_cv():
    # Creates a new Word document
    doc = Document()

    # Adds name and contact information
    doc.add_heading('SHALOM GAD', level=1)
    doc.add_paragraph(f'🌐 Website: [https://ewaesaa.com]')
    doc.add_paragraph(f'📦 GitHub: [https://github.com/gadcode]')
    doc.add_paragraph(f'📧 E-mail: [shalomgad@proton.me]')
    doc.add_paragraph(f'📜 Certificates: [https://drive.proton.me/urls/96MD0BJ03R#wu0zvy0FCTmx]')
    doc.add_paragraph('📍 Address: 7 Ozara Main Street, Lagos, Nigeria')

    # Adds introduction
    doc.add_heading('DISTINGUISHED SECURITY ENGINEER', level=2)
    doc.add_paragraph(
        "Threat Modeling | Cloud Computing | Security Hardening | Incident Detection and Response | Malware Prevention | Computer Forensics\n"
        "An exceptionally motivated cybersecurity professional with a strong foundation in software development "
        "and a passion for securing digital environments. Eager to leverage technical expertise, strategic thinking, and "
        "and hands-on experience to protect organizations from cyber threats and ensure a robust security posture."
    )

    # Adds key skills
    doc.add_heading('Key Skills', level=2)
    skills = [
        "Cyber Security Fundamentals",
        "Threat Detection and Mitigation",
        "Security Solutions Design",
        "Incident Response & Forensics",
        "Network & Cloud Security",
        "Vulnerability Assessment",
        "Security Awareness & Training"
    ]
    for skill in skills:
        doc.add_paragraph(f"- {skill}")

    # Adds Certifications
    doc.add_heading('Certifications', level=2)
    certifications = [
        "Google Cybersecurity Professional",
        "ISC2 Official Certified in Cybersecurity Course Completion"
    ]
    for cert in certifications:
        doc.add_paragraph(f"- {cert}")

    # Adds Education
    doc.add_heading('Education', level=2)
    doc.add_paragraph(
        "- Bachelor of Science in Business Administration (In Progress)\n"
        "  University of the People, Pasadena, California"
    )

    # Adds key accomplishments
    doc.add_heading('Key Accomplishments', level=2)
    accomplishments = [
        "- Authored fun-loving standard operating procedures (SOPs) which improved safety 15% and licensed 20% efficiency",
        "- Saved $200,000 from operating cost with state-of-the-art methodologies for consistent stimulation of production processes.",
        "- Chaired and supervised continuous process improvement that stimulated 30% productivity and eliminated wastes 85%.",
        "- Brainstormed business re-engineering that incremented productivity 16%, reduced labor cost, and decreased inventory by 80%.",
        "- Deployed process automation that cut 10 hours of work, saving $15,000 a month—honored with extensive training & development.",
        "- Disarmed antagonistic disadvantages; tripling quality from 12% to 39%, and leading company to its first most profitable quarter ever.",
        "- Exceeded profitability goals 14%, through addressed, remodeled, and rock-solidified quality assurance (QA)/quality control (QC)."
    ]
    for accomplishment in accomplishments:
        doc.add_paragraph(accomplishment)

    # Adds Additional Information
    doc.add_heading('Core Skills', level=2)
    additional_skills = [
        "Security Tools: Wireshark, Metasploit, tcpdump, Snort",
        "Programming Languages: Python, SQL, JavaScript, C, Linux Shell Scripting",
        "Cybersecurity Concepts: Encryption, Vulnerability Assessment, Identity and Access Management",
        "Cloud Security: AWS, GCP, Microsoft Azure",
        "Incident Response: Analysis and Resolution",
        "Regulatory (GDPR, PCI DSS, HIPAA) Compliance and Risk Management"
    ]
    for skill in additional_skills:
        doc.add_paragraph(f"- {skill}")

    # Adds community involvement and interests
    doc.add_heading('Community Involvement and Interests', level=2)
    community_interests = [
        "- Volunteer at local cybersecurity workshops and initiatives.",
        "- Ethical Hacking enthusiast",
        "- Fluent in English; keen to learn German"
        "- Passion for family time, sustainability, and philanthropy."
    ]
    for interest in community_interests:
        doc.add_paragraph(interest)

    # Saves rhe document
    doc.save('distinguished_security_engineer.docx')

if __name__ == "__main__":
    create_cv()
    print("CV successfully created as distinguished_security_engineer.docx")
